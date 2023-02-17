import docxtpl, docx
import math
from docx.shared import Inches
from docxcompose.composer import Composer


class ModelReport:
    """ Super class for all the classes that would handle rendering timetable data on the .docx file.
    institution, director, session, acronym, extra_info will be dynamically added at runtime

    It works like this:
    1. First get the template_file as a Document() object
    2. Do all the page adjustment gymnastics on it.
    3. Save it into one of the folders of the app.
    4. Apply Docxtpl on it


    THE SAME FILENAME IS USED FOR THE DOCX-ADJSUTED TEMPLATE FILE, SO IT GETS OVERRIDDEN EVERYTIME!!   """
    def __init__(self, template_name, output_filename):
        self.template_file = template_name

        self.document = docx.Document(self.template_file)
        # self._doc = docxtpl.DocxTemplate(self.template_file)

        self.sections = self.document.sections
        self.output_filename = output_filename if output_filename.endswith((".docx",".doc")) else output_filename + ".docx"
        self.context = {}
        self.template_orientation = "portrait"
        self.institution=None
        self.director=None
        self.session=None
        self.acronym=None
        self.extra_info=None


    def _calc_page_size(self, column_lower_bound, column_upper_bound, paper_size):
        """ This function calculates the width and height of the docx page based on the length ofthe longest period.
        This function returns the paper-size, as well as an INSTRUCTIONS as to the orientation it should be set to. """
        
        # Returns the length of the longest period today (+ 1 for particulars or something)
        width_limit = self._get_max_period_length() + 1
        # The width and height of an A4 document
        if width_limit <= column_lower_bound:
            paper_size = map(Inches, paper_size)
            return paper_size, "P"

        scale_factor = math.ceil((width_limit - column_lower_bound) / (column_upper_bound - column_lower_bound))
        paper_width, paper_height = Inches(round(scale_factor * paper_size[0], 2)), Inches(round(scale_factor * paper_size[1], 2))
        return (paper_width, paper_height), "L"


    def set_size_and_orientation(self, column_lower_bound=6, column_upper_bound=14, paper_size=(8.27, 11.69)):
        """ Sets the size and orientation of a page accoriding to the longest paeriods available """
        (width, height), orient = self._calc_page_size(column_lower_bound, column_upper_bound, paper_size)

        # Change the orientation to portrait of landscape depending...
        size = (width, height) if orient == "P" else (height, width)
        self.sections[0].page_width, self.sections[0].page_height = size


    def save_adjusted_template_file(self):
        """ Handles saving the readjusted document template file """
        self.adj_file_name = "adjusted_template_file.docx"
        self.document.save(self.adj_file_name)


    def begin_template_rendering(self):
        """ This method runs the  newly adjusted file with self.adj_file_name and loads it as a docxtpl template """
        self.adjusted_doc = docxtpl.DocxTemplate(self.adj_file_name)


    def populate_context(self):
        """ Populates the context dictioanry with some values that will be available to all the subclasses """
        self.context["institution"] = self.institution
        self.context["director"] = self.director
        self.context["session"] = self.session
        self.context["acronym"] = self.acronym
        self.context["extra_info"] = self.extra_info


    def run(self):
        """ Renders the template into the ms-word document """

        # ---- Integrate the needed methods to readjust this file should the need arise
        self.set_size_and_orientation()
        self.save_adjusted_template_file()
        self.begin_template_rendering()

        self.populate_context()
        self.adjusted_doc.render(self.context)
        self.adjusted_doc.save(self.output_filename)


# ----------------------------------------------------------------------------------------------------------------
# -----------------------------------------------------------------------------------------------------------------------------
class ReportBySchoolClass(ModelReport):
    """ Generates report by on a School class category - School class - school class arm basis """
    def __init__(self, classgroup, template_name, output_filename="/report_by_class"):
        super().__init__(template_name, output_filename)
        self.classgroup = classgroup


    def _get_max_period_length(self):
        """ Returns the max width ofthe periods of all the class arms in every class of this class group """    
        return max([len(periods_list) for clss in self.classgroup.school_class_list for
         arm in clss.school_class_arm_list for periods_list in arm.periods.values()])
        

    def populate_context(self):
        """ populates the context dictionary with school class and school class arm data to be used to generate templates """
        super().populate_context()
        self.context["classgroup_name"] = self.classgroup.tag()
        self.context["classgroup_fullname_id"] = self.classgroup.full_name
        self.context["classgroup_position"] = self.classgroup.id
        self.context["sch_class_list"] = [{"sch_class_name":sch_class.sch_class_alias, "sch_class_fullname":sch_class.full_name,
                "arms_list": [{"arm_name": arm.full_name, "periods_per_day":[{"day":day.day, "periods":[period.period_name for period in periods_list]}
                for day, periods_list in arm.periods.items()]} for arm in sch_class.school_class_arm_list]} 
                for sch_class in self.classgroup.school_class_list]



class ReportByDay(ModelReport):
    """ Generates Ms-Word report of all the class arms and the subjects they offer on a day-by-day basis.
    day is an object, not a string """
    def __init__(self, day, template_name, output_filename="/report_by_day"):
        super().__init__(template_name, output_filename)
        self.day = day

    @property
    def class_groups(self):
        return self.day.get_classgroups_today()

    def populate_context(self):
        """ populates the context dictionary which will be used on the .docx template """
        super().populate_context()

        self.context["day_name"] = self.day.day
        self.context["day_position"] = self.day.id
        self.context["day_fullname_id"] = self.day.full_name
        self.context["class_groups"] = [{"classgroup_name":classgroup.full_name, "description":classgroup.description,
        "sch_classes_list":[{"name":sch_class.full_name,"arms_list":[{"name": arm.full_name, "periods": [period.period_name
                for period in self.day.get_classarm_periods_for_today(arm)]} for arm in self.day.get_classarms_from_sch_classes(sch_class)]

            } for sch_class in self.day.get_sch_classes_from_class_group_today(classgroup)]
        } for classgroup in self.class_groups]



class ReportByFaculty(ModelReport):
    """ generates ms-word reports on departments(faculty) - subject - teacher - basis. <template_filename>, <output_filename>
    are strings of the path into which the template is to be generated. dept_obj is the faculty (or department) object to be generated.
    faculty is rendered as DEPARTMENT in the code. bear with me """
    def __init__(self, faculty, template_name, output_filename="/report_by_departments"):
        super().__init__(template_name, output_filename)
        self.faculty = faculty
        # self.output_filename = output_filename if output_filename.endswith((".docx",".doc")) else output_filename + ".docx"

    def populate_context(self):
        """ populates the context dictionary so it can be rendered in the word template.
        the subjects are definitely academic periods """
        super().populate_context()
        self.context["department_name"] = self.faculty.name
        self.context["department_fullname_id"] = self.faculty.full_name
        self.context["department_description"] = self.faculty.description
        self.context["department_hod"] = self.faculty.HOD
        self.context["department_subjects"] = [
        {
            "name":dept.dept_name, "hos":dept.hos,
            "teachers":[
            {"fullname": teacher.full_name, "is_exclusive":teacher.is_exclusive(),
            "other_subjects": teacher.other_depts_taught_str, "classarms_taught":teacher.classarms_taught_based_on_dept_str(dept),
            "specialization":teacher.specialization, "designation":teacher.designation} for teacher in dept.teachers_list
        ]
        } for dept in self.faculty.depts_list]



class Reporter:
    """ Container class to supply the necessary faculty, dept or day data to the classes that generate their reports """
    def __init__(self, tt_obj):
        self.tt_obj = tt_obj


    def merge_into_one_file(self, files_list, file_title="Report_on_models"):
        """ Haanles merging all the files in 'files_list' into one file with filename "master_title". """
        master_filename = file_title if file_title.endswith((".docx",".doc")) else file_title + ".docx"
        master = docx.Document(files_list[0])
        composer = Composer(master)
        
        for file in files_list[1:]:
            doc_temp = docx.Document(file)
            composer.append(doc_temp)
        composer.save(master_filename)


    def generate_report_by_faculty(self, file_title="Report_By_Faculty"):        
        """ Handles the generation of reports for all the faculties available in the timetable """

        faculty_files = []
        for fac_obj in self.tt_obj.list_of_faculties:
            report = ReportByFaculty(fac_obj, "TIMETABLE/reports/report_templates/render_models/departments_template.docx", 
            output_filename=f"{fac_obj.full_name}.docx")

            report.institution = self.tt_obj.institution
            report.director = self.tt_obj.director
            report.session = self.tt_obj.session_or_year
            report.acronym = self.tt_obj.acronym
            report.extra_info = self.tt_obj.extra_info

            report.run()
            faculty_files.append(f"{fac_obj.full_name}.docx")

        # Merge files into one
        self.merge_into_one_file(faculty_files, file_title=file_title)


    def generate_report_by_day(self, file_title="Report_By_Day"):
        """ Handles the generation of reports for all the days available in the timetable. """
        days_files = []
        for day_obj in self.tt_obj.list_of_days:
            report = ReportByDay(day_obj, "TIMETABLE/reports/report_templates/render_models/day_template.docx",
            output_filename=f"{day_obj.full_name}.docx")

            report.institution = self.tt_obj.institution
            report.director = self.tt_obj.director
            report.session = self.tt_obj.session_or_year
            report.acronym = self.tt_obj.acronym
            report.extra_info = self.tt_obj.extra_info

            report.run()
            days_files.append(f"{day_obj.full_name}.docx")

        # Merge files into one
        self.merge_into_one_file(days_files, file_title=file_title)



    def generate_report_by_school_class(self, file_title="Report_By_Class_categories"):
        """ Handles the generation of reports for all the classgroups, classes and arms available in the timetable """
        classgroup_files = []
        for classgroup in self.tt_obj.list_of_school_class_groups:
            report = ReportBySchoolClass(classgroup, "TIMETABLE/reports/report_templates/render_models/class_template.docx",
            output_filename=f"{classgroup.full_name}.docx")

            report.institution = self.tt_obj.institution
            report.director = self.tt_obj.director
            report.session = self.tt_obj.session_or_year
            report.acronym = self.tt_obj.acronym
            report.extra_info = self.tt_obj.extra_info

            report.run()
            classgroup_files.append(f"{classgroup.full_name}.docx")

        # Merge files into one
        self.merge_into_one_file(classgroup_files, file_title=file_title)

